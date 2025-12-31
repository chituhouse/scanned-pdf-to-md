#!/usr/bin/env python3
"""
生成带目录索引的Word文档
"""

import json
import re
import os
from datetime import datetime

from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# 路径配置
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
PROJECT_DIR = os.path.dirname(SCRIPT_DIR)
INPUT_FILE = os.path.join(PROJECT_DIR, "output/processed/questions_final.json")
OUTPUT_FILE = os.path.join(PROJECT_DIR, "output/公共营养师三级历年真题.docx")


def add_toc(doc):
    """插入目录字段（需要在Word中刷新）"""
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()

    # 创建TOC字段
    fld_char_begin = OxmlElement('w:fldChar')
    fld_char_begin.set(qn('w:fldCharType'), 'begin')

    instr_text = OxmlElement('w:instrText')
    instr_text.set(qn('xml:space'), 'preserve')
    instr_text.text = 'TOC \\o "1-1" \\h \\z \\u'  # 只索引Heading 1

    fld_char_separate = OxmlElement('w:fldChar')
    fld_char_separate.set(qn('w:fldCharType'), 'separate')

    fld_char_end = OxmlElement('w:fldChar')
    fld_char_end.set(qn('w:fldCharType'), 'end')

    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_separate)
    run._r.append(fld_char_end)

    # 添加提示文字
    doc.add_paragraph('（请右键点击此处，选择"更新域"以生成目录）')
    doc.add_paragraph()


def set_chinese_font(run, font_name='宋体', font_size=12):
    """设置中文字体"""
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)


def parse_markdown_table(text):
    """
    解析Markdown表格，返回表格数据
    Returns: [(row1_cells), (row2_cells), ...] 或 None
    """
    lines = text.strip().split('\n')
    if len(lines) < 2:
        return None

    tables = []
    current_table = []

    for line in lines:
        if line.strip().startswith('|'):
            # 跳过分隔行
            if re.match(r'^\|[\s\-:|]+\|$', line.strip()):
                continue
            # 解析单元格
            cells = [cell.strip() for cell in line.strip().split('|')[1:-1]]
            if cells:
                current_table.append(cells)
        else:
            if current_table and len(current_table) >= 2:
                tables.append(current_table)
            current_table = []

    if current_table and len(current_table) >= 2:
        tables.append(current_table)

    return tables if tables else None


def add_table_to_doc(doc, table_data):
    """将表格数据添加到Word文档"""
    if not table_data or len(table_data) < 1:
        return

    # 获取列数（取最大列数）
    max_cols = max(len(row) for row in table_data)

    # 创建表格
    table = doc.add_table(rows=len(table_data), cols=max_cols)
    table.style = 'Table Grid'

    for i, row_data in enumerate(table_data):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            if j < max_cols:
                cell = row.cells[j]
                cell.text = cell_text
                # 设置字体
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        set_chinese_font(run, font_size=10)

    doc.add_paragraph()  # 表格后空行


def is_toc_page(text):
    """判断是否为目录页"""
    if '目录' in text[:50]:
        return True
    page_refs = re.findall(r'真题(?:答案)?\s*\d{2,3}', text)
    return len(page_refs) > 5


def is_answer_section(text):
    """判断是否为答案部分"""
    first_150 = text[:150]
    return '答案' in first_150 or '解析' in first_150


def add_content_to_doc(doc, content, is_table_page=False):
    """添加页面内容到文档"""
    if not content:
        return

    # 检查是否包含Markdown表格
    if '|' in content and is_table_page:
        lines = content.split('\n')
        non_table_lines = []
        table_buffer = []
        in_table = False

        for line in lines:
            if line.strip().startswith('|'):
                if not in_table:
                    # 先输出之前的非表格内容
                    if non_table_lines:
                        text = '\n'.join(non_table_lines).strip()
                        if text:
                            p = doc.add_paragraph(text)
                            for run in p.runs:
                                set_chinese_font(run)
                        non_table_lines = []
                in_table = True
                table_buffer.append(line)
            else:
                if in_table:
                    # 表格结束，解析并添加
                    tables = parse_markdown_table('\n'.join(table_buffer))
                    if tables:
                        for table_data in tables:
                            add_table_to_doc(doc, table_data)
                    table_buffer = []
                    in_table = False
                non_table_lines.append(line)

        # 处理末尾
        if table_buffer:
            tables = parse_markdown_table('\n'.join(table_buffer))
            if tables:
                for table_data in tables:
                    add_table_to_doc(doc, table_data)

        if non_table_lines:
            text = '\n'.join(non_table_lines).strip()
            if text:
                p = doc.add_paragraph(text)
                for run in p.runs:
                    set_chinese_font(run)
    else:
        # 普通文本
        # 按段落分割，保持格式
        paragraphs = content.split('\n\n')
        for para_text in paragraphs:
            para_text = para_text.strip()
            if para_text:
                p = doc.add_paragraph(para_text.replace('\n', ' '))
                for run in p.runs:
                    set_chinese_font(run)


def generate_word():
    """生成Word文档"""
    print("读取数据...")
    with open(INPUT_FILE, 'r', encoding='utf-8') as f:
        data = json.load(f)

    pages = {p['page_num']: p for p in data['pages']}
    exams = data['exams']

    print(f"共 {len(pages)} 页，{len(exams)} 套考试")

    # 创建文档
    doc = Document()

    # 设置页面边距
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # 添加标题
    title = doc.add_heading('公共营养师三级历年真题及答案解析', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 添加文档信息
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M')}")
    set_chinese_font(run, font_size=10)

    doc.add_paragraph()

    # 添加目录
    toc_title = doc.add_heading('目录', 1)
    add_toc(doc)

    # 添加分页符
    doc.add_page_break()

    # 整理考试数据
    exam_list = []
    for exam in exams:
        start_page = exam['start_page']
        first_page = pages.get(start_page, {})
        content = first_page.get('markdown', '') or '\n'.join(first_page.get('text', []))

        # 跳过目录页
        if is_toc_page(content):
            continue

        year, month = exam['exam_id'].split('-')
        is_answer = is_answer_section(content)

        exam_list.append({
            'exam_id': exam['exam_id'],
            'year': year,
            'month': month,
            'is_answer': is_answer,
            'start_page': start_page,
            'end_page': exam['pages'][-1],
            'pages': exam['pages'],
            'title': f"{year}年{int(month)}月公共营养师三级" + ("答案解析" if is_answer else "统考真题")
        })

    # 按年月排序
    exam_list.sort(key=lambda x: (x['year'], x['month'], x['is_answer']), reverse=True)

    print(f"处理 {len(exam_list)} 套考试/答案...")

    # 正文内容
    for idx, exam in enumerate(exam_list):
        print(f"  [{idx+1}/{len(exam_list)}] {exam['title']}")

        # 添加考试标题（Heading 1，会被目录索引）
        doc.add_heading(exam['title'], 1)

        for page_num in exam['pages']:
            page = pages.get(page_num, {})
            content = page.get('markdown', '') or '\n'.join(page.get('text', []))

            if not content.strip():
                continue

            # 跳过目录页内容
            if is_toc_page(content):
                continue

            # 添加内容
            is_table_page = page.get('is_table_page', False)
            add_content_to_doc(doc, content, is_table_page)

        # 每套考试后添加分页符（除了最后一套）
        if idx < len(exam_list) - 1:
            doc.add_page_break()

    # 保存文档
    doc.save(OUTPUT_FILE)

    file_size = os.path.getsize(OUTPUT_FILE) / 1024
    print(f"\nWord文档已生成: {OUTPUT_FILE}")
    print(f"文件大小: {file_size:.1f} KB")
    print(f"包含 {len(exam_list)} 套考试/答案")
    print('\n提示：打开Word后，右键点击目录区域，选择"更新域"即可生成完整目录索引')


if __name__ == "__main__":
    generate_word()
